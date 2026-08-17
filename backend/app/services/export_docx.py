"""Markdown → Word (.docx) export for research papers."""

from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


def markdown_to_docx_bytes(title: str, content_md: str) -> bytes:
    """Convert research markdown into a downloadable Word document."""
    doc = Document()
    _set_default_style(doc)

    doc.add_heading(title or "Research Export", level=0)

    text = (content_md or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines) if code_lines else "")
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Markdown table (header | --- | body)
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+:?\s*\|", lines[i + 1]):
            table_lines = [stripped]
            i += 1
            # skip separator
            i += 1
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1
            _add_table(doc, table_lines)
            continue

        if stripped in {"---", "***", "___"}:
            _add_horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(stripped[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # Ensure East Asian / complex scripts also get Calibri when Word substitutes
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level, size in ((1, 16), (2, 13), (3, 12)):
        try:
            hs = doc.styles[f"Heading {level}"]
            hs.font.name = "Calibri"
            hs.font.size = Pt(size)
            hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        except KeyError:
            pass


def _split_table_row(line: str) -> list[str]:
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def _add_table(doc: Document, table_lines: list[str]) -> None:
    if not table_lines:
        return
    rows = [_split_table_row(line) for line in table_lines]
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "666666")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_inline_runs(paragraph, text: str) -> None:
    # Inline markdown: **bold**, *italic*, `code`, [label](url)
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|[^*`\[]+|\*+|[^\*]+)"
    )
    for token in pattern.findall(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif token.startswith("[") and "](" in token and token.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if m:
                label, url = m.group(1), m.group(2)
                run = paragraph.add_run(f"{label} ({url})")
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            else:
                paragraph.add_run(token)
        else:
            paragraph.add_run(token)
